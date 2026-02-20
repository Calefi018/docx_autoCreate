import streamlit as st
from docx import Document
import google.generativeai as genai
import io

# Puxa a chave de forma segura
CHAVE_API = st.secrets["GEMINI_API_KEY"]
genai.configure(api_key=CHAVE_API)

# ---------------------------------------------------------
# FUNÇÕES DE LEITURA E ESCRITA DE WORD
# ---------------------------------------------------------
def extrair_texto_docx(arquivo_upload):
    """Lê o arquivo Word enviado pelo usuário e extrai todo o texto."""
    doc = Document(arquivo_upload)
    texto_completo = []
    for paragrafo in doc.paragraphs:
        if paragrafo.text.strip():
            texto_completo.append(paragrafo.text)
    return "\n".join(texto_completo)

def criar_gabarito_word(texto_ia):
    """Cria um novo documento Word formatando o texto gerado pela IA (negritos e parágrafos)."""
    doc = Document()
    doc.add_heading('Gabarito Gerado - Desafio Profissional', level=1)
    
    linhas = texto_ia.split('\n')
    for linha in linhas:
        linha = linha.strip()
        if not linha:
            continue
            
        # Lógica simples para aplicar negrito onde a IA usou **texto**
        p = doc.add_paragraph()
        partes = linha.split('**')
        
        for i, parte in enumerate(partes):
            if i % 2 == 1: # Ímpares são os textos que estavam entre ** **
                p.add_run(parte).bold = True
            else:
                p.add_run(parte)
                
    arquivo_saida = io.BytesIO()
    doc.save(arquivo_saida)
    return arquivo_saida.getvalue()

# ---------------------------------------------------------
# FUNÇÃO DE IA PARA LER TEMPLATE + TEMA E GERAR RESPOSTA
# ---------------------------------------------------------
def gerar_resolucao_inteligente(texto_template, texto_tema, nome_modelo):
    """Envia o template E o tema descritivo para a IA interpretar e resolver."""
    modelo = genai.GenerativeModel(nome_modelo)
    
    prompt = f"""
    Você é um especialista acadêmico ajudando um estudante universitário a resolver um Desafio Profissional.
    
    Vou te fornecer duas informações cruciais:
    1. A DESCRIÇÃO DO TEMA/CASO que precisa ser resolvido.
    2. O TEXTO DO TEMPLATE PADRÃO ÚNICO, que dita a estrutura das Etapas (1 a 5) que você deve preencher.
    
    Sua tarefa é gerar todas as respostas passo a passo, de forma original e sem plágio, resolvendo o caso apresentado.
    Me informe claramente onde preencher as informações no Word (Ex: "Na Etapa 2 (Materiais de referência), escreva isso...").
    
    REGRA OBRIGATÓRIA PARA A ETAPA 5 (MEMORIAL ANALÍTICO):
    Você deve redigir a Etapa 5 seguindo estritamente o padrão abaixo. Escreva os títulos em **negrito**:
    
    **Resumo:** [Escreva 1 parágrafo resumindo o que descobriu]
    **Contextualização do desafio:** [Escreva 1 parágrafo: Quem? Onde? Qual a situação?]
    **Análise:** [Escreva 1 parágrafo usando de 2 a 3 conceitos da disciplina, com exemplos]
    **Propostas de solução:** [Escreva de 1 a 2 parágrafos com recomendações e teorias que as apoiam]
    **Conclusão reflexiva:** [Escreva de 1 a 2 parágrafos sobre o que foi aprendido]
    **Referências:** [Liste as referências em formato ABNT, incluindo o livro base da disciplina se mencionado no tema]
    **Autoavaliação:** [Escreva 1 parágrafo sobre o processo de estudo]
    
    Checklist que você deve respeitar:
    - O texto total da Etapa 5 NÃO pode passar de 6000 caracteres.
    - Os conceitos devem fazer sentido e conectar teoria e situação.
    - Apresentar soluções plausíveis.
    
    =========================================
    DESCRIÇÃO DO TEMA/CASO (O problema a ser resolvido):
    {texto_tema}
    
    =========================================
    TEXTO DO TEMPLATE (A estrutura a ser preenchida):
    {texto_template}
    """
    
    try:
        resposta = modelo.generate_content(prompt)
        return resposta.text
    except Exception as e:
        st.error(f"Erro da IA ({nome_modelo}): {e}")
        return None

# ---------------------------------------------------------
# INTERFACE DO SITE (STREAMLIT)
# ---------------------------------------------------------
st.set_page_config(page_title="Gerador Automático de Trabalhos", page_icon="🎓")

st.title("Gerador Universal - Desafio Profissional 🎓")
st.write("Anexe o template padrão e cole as instruções do portal. A IA fará o resto!")

# Variáveis de estado
if "gabarito_pronto" not in st.session_state:
    st.session_state.gabarito_pronto = None
if "texto_tela" not in st.session_state:
    st.session_state.texto_tela = ""

# Lista modelos
modelos_disponiveis = []
try:
    for m in genai.list_models():
        if 'generateContent' in m.supported_generation_methods:
            modelos_disponiveis.append(m.name.replace('models/', ''))
except Exception as e:
    st.error("Erro ao conectar com a API do Google. Verifique sua chave.")

if modelos_disponiveis:
    modelo_escolhido = st.selectbox("Selecione a IA:", modelos_disponiveis, index=0)
    
    # 1. Upload do Arquivo
    arquivo_upload = st.file_uploader("1. Faça o upload do Template Padrão (.docx)", type=["docx"])
    
    # 2. Caixa de texto para o Tema
    tema_desafio = st.text_area(
        "2. Cole aqui o Tema/Descrição do Desafio (copiado do portal)", 
        height=250, 
        placeholder="Cole aqui o texto inteiro do desafio (ex: Você é responsável por coordenar a implantação de tecnologias emergentes em uma obra...)"
    )

    if st.button("Analisar e Gerar Trabalho", type="primary"):
        if arquivo_upload is not None and tema_desafio.strip() != "":
            with st.spinner(f"Lendo o template e resolvendo o caso com {modelo_escolhido}... isso pode levar alguns segundos."):
                
                # Extrai o texto do Word
                texto_do_template = extrair_texto_docx(arquivo_upload)
                
                # Manda pra IA resolver juntando as duas coisas
                resposta_ia = gerar_resolucao_inteligente(texto_do_template, tema_desafio, modelo_escolhido)
                
                if resposta_ia:
                    # Salva para exibir na tela e gera o Word de Gabarito
                    st.session_state.texto_tela = resposta_ia
                    st.session_state.gabarito_pronto = criar_gabarito_word(resposta_ia)
                    st.success("✅ Trabalho gerado com sucesso!")
        else:
            st.warning("⚠️ Por favor, faça o upload do template Word E cole o tema do desafio antes de gerar.")
            
    # Mostra o resultado e o botão de download
    if st.session_state.gabarito_pronto is not None:
        st.download_button(
            label="⬇️ Baixar Gabarito em Word (.docx)",
            data=st.session_state.gabarito_pronto,
            file_name="Gabarito_Desafio_Profissional.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        st.markdown("---")
        st.subheader("Pré-visualização do Resultado:")
        st.markdown(st.session_state.texto_tela)

else:
    st.error("Nenhum modelo compatível encontrado.")
